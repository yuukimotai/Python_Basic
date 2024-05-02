import docx

doc = docx.Document('word_sample.docx')
doc.add_heading('Document Title', 0)
p = doc.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic').italic = True
doc.save('word_sample.docx')